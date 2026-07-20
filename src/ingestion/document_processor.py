from pathlib import Path
from src.common.logging import get_logger
from src.common.exceptions import ValidationError

logger = get_logger(__name__)


class ProcessedDocument:
    def __init__(self, text: str, metadata: dict):
        self.text = text
        self.metadata = metadata


def process_document(file_path: Path, original_name: str) -> ProcessedDocument:
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return _process_pdf(file_path, original_name)
    elif ext == ".docx":
        return _process_docx(file_path, original_name)
    elif ext == ".odt":
        return _process_odt(file_path, original_name)
    elif ext == ".txt":
        return _process_txt(file_path, original_name)
    else:
        raise ValidationError(f"Unsupported file format: {ext}")


def _process_pdf(file_path: Path, original_name: str) -> ProcessedDocument:
    try:
        import pypdf
    except ImportError:
        raise ValidationError("pypdf is required for PDF processing")
    reader = pypdf.PdfReader(file_path)
    text_parts: list[str] = []
    metadata = {
        "title": original_name,
        "pages": len(reader.pages),
        "format": "pdf",
    }
    if reader.metadata:
        if reader.metadata.title:
            metadata["title"] = reader.metadata.title
        if reader.metadata.author:
            metadata["author"] = reader.metadata.author
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        text_parts.append(f"--- Página {i + 1} ---\n{page_text}")
    return ProcessedDocument(text="\n\n".join(text_parts), metadata=metadata)


def _process_docx(file_path: Path, original_name: str) -> ProcessedDocument:
    try:
        from docx import Document
    except ImportError:
        raise ValidationError("python-docx is required for DOCX processing")
    doc = Document(file_path)
    text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
    metadata = {
        "title": original_name,
        "format": "docx",
    }
    return ProcessedDocument(text="\n\n".join(text_parts), metadata=metadata)


def _process_odt(file_path: Path, original_name: str) -> ProcessedDocument:
    try:
        from odf import text, teletype
        from odf.opendocument import load
    except ImportError:
        raise ValidationError("odfpy is required for ODT processing")
    doc = load(file_path)
    text_parts: list[str] = []
    for paragraph in doc.getElementsByType(text.P):
        text_parts.append(teletype.extractText(paragraph))
    metadata = {
        "title": original_name,
        "format": "odt",
    }
    return ProcessedDocument(text="\n\n".join(text_parts), metadata=metadata)


def _process_txt(file_path: Path, original_name: str) -> ProcessedDocument:
    text = file_path.read_text(encoding="utf-8", errors="replace")
    metadata = {
        "title": original_name,
        "format": "txt",
    }
    return ProcessedDocument(text=text, metadata=metadata)
